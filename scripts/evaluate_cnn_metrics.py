#!/usr/bin/env python3
"""
Offline & Online Evaluation Script for CIL Waypoint Model
Calculates: ADE, FDE, FPS, Latency, Route Complete, TTC

Usage:
  # Full offline evaluation (requires H5 dataset):
  set H5_PATH=C:/path/to/carla_images_drive.h5
  python scripts/evaluate_cnn_metrics.py

  # Benchmark-only mode (no dataset needed, measures FPS/Latency):
  python scripts/evaluate_cnn_metrics.py --benchmark-only

  # Show online metrics from last CARLA run:
  python scripts/evaluate_cnn_metrics.py --online-log outputs/telemetry.csv
"""

import argparse
import csv
import os
import sys
import tempfile
import time
from pathlib import Path

import torch
import numpy as np

# python-docx (pip install python-docx) — only required for docx export
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# Add project root to path
PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from core_perception.cnn_model import WaypointPredictor


# ============================================================================
# CSV UTILITIES
# ============================================================================
def split_combined_csv_to_towns(combined_csv: str, out_dir: Path) -> Path:
    """
    Split a single combined driving_log.csv (all towns in one file) into
    per-town CSV files expected by WaypointCarlaDatasetH5 / find_csv_root:
        <out_dir>/Town01/driving_log.csv
        <out_dir>/Town02/driving_log.csv
        ...
    Returns out_dir (the csv_root to pass downstream).
    """
    town_rows: dict = {}
    fieldnames = None

    with open(combined_csv, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        for row in reader:
            # Infer town from center_camera path (e.g. "Town01/images_center/3.jpg")
            center = row.get("center_camera", "")
            parts = center.replace("\\", "/").split("/")
            town = parts[0] if parts and parts[0].startswith("Town") else None
            if town is None:
                continue
            town_rows.setdefault(town, []).append(row)

    if not town_rows:
        raise ValueError(
            f"No Town*/... paths found in {combined_csv}. Check center_camera column."
        )

    out_dir.mkdir(parents=True, exist_ok=True)
    for town, rows in town_rows.items():
        town_dir = out_dir / town
        town_dir.mkdir(parents=True, exist_ok=True)
        with open(town_dir / "driving_log.csv", "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)
        print(f"  [CSV] {town}: {len(rows)} rows -> {town_dir / 'driving_log.csv'}")

    print(f"  [CSV] Split into {len(town_rows)} town CSVs under: {out_dir}")
    return out_dir


# ============================================================================
# MODEL DISCOVERY
# ============================================================================
def _remap_film_keys(state_dict: dict) -> tuple[dict, bool]:
    """
    Remap old single-FiLM checkpoint keys (film.*) to multi-stage keys (film_s4.*).
    Matches the same migration done in run_agents.py for waypoint_predictor_h5.pth.
    film_s3 is left as identity (not in checkpoint) so the model uses default init.
    Returns (remapped_state_dict, did_remap).
    """
    remapped = {}
    did_remap = False
    for k, v in state_dict.items():
        # Old key: "film.gamma", "film.beta", etc.
        # New key: "film_s4.gamma", "film_s4.beta", etc.
        if k.startswith("film.") and not any(
            k.startswith(prefix) for prefix in ("film_s1.", "film_s2.", "film_s3.", "film_s4.")
        ):
            new_key = "film_s4." + k[len("film."):]
            remapped[new_key] = v
            did_remap = True
        else:
            remapped[k] = v
    return remapped, did_remap


def find_best_checkpoint(models_dir: Path) -> Path:
    """Auto-detect the best waypoint model checkpoint."""
    # Priority order
    candidates = [
        "waypoint_predictor.pth",
        "waypoint_predictor_h5.pth",
    ]
    for name in candidates:
        p = models_dir / name
        if p.exists():
            return p

    # Fallback: any file with 'waypoint' or 'cil' in name
    for pattern in ["*waypoint*.pth", "*cil*.pth"]:
        matches = sorted(models_dir.glob(pattern), key=lambda x: x.stat().st_mtime, reverse=True)
        if matches:
            return matches[0]

    return models_dir / "waypoint_predictor.pth"  # will fail later with clear message


def load_model(device: torch.device) -> torch.nn.Module:
    """Load WaypointPredictor model with auto-detected checkpoint."""
    models_dir = PROJECT_ROOT / "models"
    checkpoint_path = find_best_checkpoint(models_dir)

    if not checkpoint_path.exists():
        print(f"[ERROR] Model checkpoint not found at: {checkpoint_path}")
        print(f"   Available .pth files in {models_dir}:")
        for f in sorted(models_dir.glob("*.pth")):
            print(f"     - {f.name} ({f.stat().st_size / 1024 / 1024:.1f} MB)")
        sys.exit(1)

    print(f"Loading checkpoint: {checkpoint_path}")
    model = WaypointPredictor().to(device)

    try:
        checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=True)
        if "model_state_dict" in checkpoint:
            state_dict = checkpoint["model_state_dict"]
        else:
            state_dict = checkpoint

        # Remap old single-FiLM keys (film.*) → multi-stage (film_s4.*)
        # This matches the same migration in run_agents.py for waypoint_predictor_h5.pth
        state_dict, did_remap = _remap_film_keys(state_dict)
        if did_remap:
            print("  [INFO] Remapped old film.* keys -> film_s4.* (single-FiLM checkpoint).")

        # Try strict load first, fallback to non-strict for older checkpoints
        try:
            model.load_state_dict(state_dict, strict=True)
        except RuntimeError:
            result = model.load_state_dict(state_dict, strict=False)
            missing = list(result.missing_keys)
            unexpected = list(result.unexpected_keys)

            # For old single-FiLM checkpoints (film.* → film_s4.*), film_s3.*
            # keys are intentionally absent. FiLM._init_identity() sets them to
            # gamma=1 / beta=0 (identity passthrough) — correct and safe to use
            # as-is without re-training. Suppress the noisy [WARN] for these.
            if did_remap:
                film_s3_absent = [k for k in missing if k.startswith("film_s3.")]
                missing = [k for k in missing if not k.startswith("film_s3.")]
                if film_s3_absent:
                    print(
                        f"  [INFO] film_s3.* ({len(film_s3_absent)} keys) absent in"
                        " checkpoint — identity init applied (gamma=1, beta=0 passthrough)."
                    )

            if missing:
                print(f"  [WARN] Missing keys (using defaults): {len(missing)} keys")
                for k in missing[:5]:
                    print(f"           {k}")
                if len(missing) > 5:
                    print(f"           ... and {len(missing) - 5} more")
            if unexpected:
                print(f"  [WARN] Unexpected keys (ignored): {len(unexpected)} keys")
                for k in unexpected[:5]:
                    print(f"           {k}")
                if len(unexpected) > 5:
                    print(f"           ... and {len(unexpected) - 5} more")
        print("[OK] Model loaded successfully.")
    except Exception as e:
        print(f"[ERROR] Error loading checkpoint: {e}")
        sys.exit(1)

    model.eval()
    return model


# ============================================================================
# BENCHMARK: FPS & LATENCY (no dataset needed)
# ============================================================================
def benchmark_inference(model: torch.nn.Module, device: torch.device,
                        num_warmup: int = 20, num_iterations: int = 200) -> dict:
    """
    Measure model inference FPS and Latency using synthetic inputs.
    Returns dict with fps, latency_ms, latency_std_ms.
    """
    print(f"\n{'-'*60}")
    print(f" [BENCH]  INFERENCE BENCHMARK (device={device})")
    print(f"{'-'*60}")

    # Create synthetic inputs matching real pipeline dimensions
    # Input: [B, 9, 66, 200] (3 temporal YUV frames, cropped & resized)
    dummy_image = torch.randn(1, 9, 66, 200, device=device)
    dummy_command = torch.tensor([0], dtype=torch.long, device=device)
    dummy_speed = torch.tensor([0.5], dtype=torch.float32, device=device)

    # Warmup (let GPU kernels compile / cache)
    print(f"  Warming up ({num_warmup} iterations)...")
    with torch.inference_mode():
        for _ in range(num_warmup):
            _ = model(dummy_image, dummy_command, dummy_speed)

    if device.type == "cuda":
        torch.cuda.synchronize()

    # Benchmark
    print(f"  Benchmarking ({num_iterations} iterations)...")
    latencies = []

    with torch.inference_mode():
        for _ in range(num_iterations):
            if device.type == "cuda":
                torch.cuda.synchronize()
            t0 = time.perf_counter()

            _ = model(dummy_image, dummy_command, dummy_speed)

            if device.type == "cuda":
                torch.cuda.synchronize()
            t1 = time.perf_counter()

            latencies.append((t1 - t0) * 1000.0)  # ms

    latencies_np = np.array(latencies)

    # Remove outliers (top/bottom 5%) for more stable measurement
    p5, p95 = np.percentile(latencies_np, [5, 95])
    filtered = latencies_np[(latencies_np >= p5) & (latencies_np <= p95)]

    mean_latency = float(np.mean(filtered))
    std_latency = float(np.std(filtered))
    fps = 1000.0 / mean_latency if mean_latency > 0 else 0.0

    # Also measure batch throughput
    batch_sizes = [1, 4, 16, 32]
    batch_fps = {}
    for bs in batch_sizes:
        try:
            b_img = torch.randn(bs, 9, 66, 200, device=device)
            b_cmd = torch.zeros(bs, dtype=torch.long, device=device)
            b_spd = torch.full((bs,), 0.5, dtype=torch.float32, device=device)

            # Warmup
            with torch.inference_mode():
                for _ in range(5):
                    _ = model(b_img, b_cmd, b_spd)

            if device.type == "cuda":
                torch.cuda.synchronize()

            t0 = time.perf_counter()
            with torch.inference_mode():
                for _ in range(50):
                    _ = model(b_img, b_cmd, b_spd)
            if device.type == "cuda":
                torch.cuda.synchronize()
            t1 = time.perf_counter()

            total_samples = 50 * bs
            batch_fps[bs] = total_samples / (t1 - t0)
            del b_img, b_cmd, b_spd
        except RuntimeError:
            batch_fps[bs] = 0.0  # OOM

    return {
        "fps": fps,
        "latency_ms": mean_latency,
        "latency_std_ms": std_latency,
        "batch_fps": batch_fps,
        "num_iterations": len(filtered),
    }


# ============================================================================
# OFFLINE: ADE & FDE (requires H5 dataset)
# ============================================================================
def evaluate_ade_fde(model: torch.nn.Module, device: torch.device,
                     h5_path: str, csv_root_hint: str | None) -> dict | None:
    """
    Calculate ADE (Average Displacement Error) and FDE (Final Displacement Error).
    Returns dict or None if dataset is unavailable.

    csv_root_hint can be:
      - A directory already structured as  <root>/Town01/driving_log.csv  (passed straight through)
      - A combined single driving_log.csv  (auto-split into per-town files in a temp dir)
      - None (find_csv_root() will try to auto-detect from h5 path)
    """
    try:
        import yaml
        import h5py
        from torch.utils.data import DataLoader
        import torchvision.transforms as transforms
        from scripts.kaggle_train_h5 import WaypointCarlaDatasetH5, find_csv_root
    except ImportError as e:
        print(f"[WARN] Cannot run ADE/FDE evaluation: {e}")
        return None

    if not os.path.exists(h5_path):
        print(f"[WARN] H5 dataset not found at: {h5_path}")
        print("   Skipping ADE/FDE evaluation.")
        print("   To enable: set H5_PATH=C:/path/to/carla_images_drive.h5")
        return None

    # Load config
    config_path = PROJECT_ROOT / "configs" / "train_params.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    print(f"\n{'-'*60}")
    print(f" [ACCURACY] OFFLINE ADE/FDE EVALUATION")
    print(f"{'-'*60}")

    try:
        # ── Resolve csv_root ──────────────────────────────────────────────────
        _tmp_dir = None  # keep alive for the lifetime of this function
        if csv_root_hint and os.path.isfile(csv_root_hint):
            # Combined single CSV: split into per-town temp structure
            print(f"  [CSV] Combined CSV detected, splitting by town...")
            _tmp_dir = tempfile.mkdtemp(prefix="cil_eval_csv_")
            # Keep as Path — WaypointCarlaDatasetH5 uses Path / operator internally
            csv_root = split_combined_csv_to_towns(csv_root_hint, Path(_tmp_dir))
        elif csv_root_hint and os.path.isdir(csv_root_hint):
            # Already a directory — wrap in Path so / operator works inside dataset
            csv_root = Path(csv_root_hint)
        else:
            raw = find_csv_root(h5_path, csv_root_hint)
            # find_csv_root may return str or Path — normalise to Path
            csv_root = Path(raw)

        with h5py.File(h5_path, "r") as f:
            towns = [k for k in f.keys() if isinstance(f[k], h5py.Group)]

        transform = transforms.Compose([transforms.Normalize(mean=[0.5] * 9, std=[0.5] * 9)])
        val_dataset = WaypointCarlaDatasetH5(
            h5_path=h5_path,
            csv_root=csv_root,
            towns=towns,
            transform=transform,
            is_training=False,
            train_ratio=float(config.get("train_split", 0.75)),
            geometric_offset=float(config.get("geometric_offset", 0.35)),
            include_side_cameras=False,
        )

        def collate_fn(batch):
            imgs, wps, cmds, recs, speeds = zip(*batch)
            return (
                torch.stack(imgs),
                torch.stack(wps).float(),
                torch.stack(cmds).long(),
                torch.tensor(recs, dtype=torch.float32),
                torch.tensor(speeds, dtype=torch.float32),
            )

        val_loader = DataLoader(
            val_dataset, batch_size=128, shuffle=False,
            num_workers=0, pin_memory=device.type == "cuda", collate_fn=collate_fn
        )
        print(f"  Validation Dataset size: {len(val_dataset)} samples")
    except Exception as e:
        print(f"[ERROR] Error preparing dataset: {e}")
        return None

    # Evaluation Loop
    print("  Evaluating...")
    total_ade = 0.0
    total_fde = 0.0
    total_samples = 0
    inference_times = []

    with torch.inference_mode():
        for i, (imgs, wps, cmds, _, speeds) in enumerate(val_loader):
            imgs = imgs.to(device, non_blocking=True)
            cmds = cmds.to(device, non_blocking=True)
            wps = wps.to(device, non_blocking=True)
            speeds = speeds.to(device, non_blocking=True)

            if device.type == "cuda":
                torch.cuda.synchronize()
            t0 = time.perf_counter()

            with torch.amp.autocast("cuda" if device.type == "cuda" else "cpu"):
                out = model(imgs, cmds, speeds)
                pred_wp = out[:, :10].view(-1, 5, 2)

            if device.type == "cuda":
                torch.cuda.synchronize()
            t1 = time.perf_counter()

            inference_times.append((t1 - t0) * 1000.0)  # ms

            # wps and pred_wp shape: [batch_size, 5, 2]
            distances = torch.norm(pred_wp - wps, dim=-1)
            batch_ade = distances.mean(dim=1).sum().item()
            batch_fde = distances[:, 4].sum().item()

            total_ade += batch_ade
            total_fde += batch_fde
            total_samples += imgs.size(0)

            if (i + 1) % 10 == 0:
                print(f"    Processed {i+1}/{len(val_loader)} batches...")

    if total_samples == 0:
        return None

    mean_ade = total_ade / total_samples
    mean_fde = total_fde / total_samples
    avg_batch_latency = float(np.mean(inference_times)) if inference_times else 0.0

    return {
        "ade_m": mean_ade,
        "fde_m": mean_fde,
        "total_samples": total_samples,
        "avg_batch_latency_ms": avg_batch_latency,
    }


# ============================================================================
# PS1 CONFIG PARSER (reads defaults from run_cil.ps1)
# ============================================================================
def _parse_ps1_defaults(ps1_path: "Path") -> dict:
    """
    Parse default parameter values from run_cil.ps1 so the online evaluation
    can mirror the exact same agent/model/CARLA configuration.

    Extracted keys: carla_root, config, agent, model, yolo_model, device.
    Falls back to sensible hard-coded defaults if parsing fails.
    """
    import re as _re
    defaults = {
        "carla_root": "E:\\Carla",
        "config":     "configs/carla_env.yaml",
        "agent":      "cil_yolo",
        "model":      "models\\waypoint_predictor_h5.pth",
        "yolo_model": "",
        "device":     "auto",
    }
    try:
        text = Path(ps1_path).read_text(encoding="utf-8")
        patterns = {
            "carla_root": r'\[string\]\$CarlaRoot\s*=\s*"([^"]+)"',
            "config":     r'\[string\]\$Config\s*=\s*"([^"]+)"',
            "agent":      r'\[string\]\$Agent\s*=\s*"([^"]+)"',
            "model":      r'\[string\]\$Model\s*=\s*"([^"]+)"',
            "yolo_model": r'\[string\]\$YoloModel\s*=\s*"([^"]+)"',
            "device":     r'\[string\]\$Device\s*=\s*"([^"]+)"',
        }
        for key, pat in patterns.items():
            m = _re.search(pat, text)
            if m:
                defaults[key] = m.group(1)
        print(f"  [PS1] Parsed run_cil.ps1: agent={defaults['agent']} | "
              f"carla_root={defaults['carla_root']} | device={defaults['device']}")
    except Exception as exc:
        print(f"  [WARN] Could not parse run_cil.ps1 defaults: {exc}")
        print("         Using built-in defaults instead.")
    return defaults


# ============================================================================
# RESULTS DISPLAY
# ============================================================================
def print_results(benchmark: dict, ade_fde: dict | None, online_metrics: dict | None = None) -> None:
    """Print comprehensive metrics report."""
    print("\n")
    print("=" * 60)
    print(" [TARGET] COMPREHENSIVE EVALUATION RESULTS")
    print("=" * 60)

    # -- Section 1: Inference Performance --
    print("\n+---------------------------------------------------------+")
    print("|  [STATS] INFERENCE PERFORMANCE                              |")
    print("+---------------------------------------------------------+")
    print(f"|  FPS (batch=1):         {benchmark['fps']:>8.1f} FPS               |")
    print(f"|  Latency (batch=1):     {benchmark['latency_ms']:>8.2f} +- {benchmark['latency_std_ms']:.2f} ms        |")
    print("|                                                         |")
    print("|  Batch Throughput:                                      |")
    for bs, fps in benchmark.get("batch_fps", {}).items():
        if fps > 0:
            per_sample = 1000.0 / fps if fps > 0 else 0
            print(f"|    batch={bs:<3d}  ->  {fps:>8.1f} FPS  ({per_sample:.2f} ms/sample)   |")
    print("+---------------------------------------------------------+")

    # -- Section 2: Trajectory Accuracy --
    print("\n+---------------------------------------------------------+")
    print("|  [ACCURACY] TRAJECTORY ACCURACY (Offline)                      |")
    print("+---------------------------------------------------------+")
    if ade_fde is not None:
        ade = ade_fde["ade_m"]
        fde = ade_fde["fde_m"]
        samples = ade_fde["total_samples"]
        print(f"|  ADE (Avg Displacement Error):  {ade:>8.4f} m             |")
        print(f"|  FDE (Final Displacement Error): {fde:>8.4f} m             |")
        print(f"|  Evaluated on: {samples:>6d} validation samples             |")

        # Performance rating
        if ade < 1.0 and fde < 1.5:
            rating = "[EXCELLENT] EXCELLENT"
        elif ade < 2.0 and fde < 3.0:
            rating = "[GOOD] GOOD"
        elif ade < 3.0 and fde < 5.0:
            rating = "[WARN] FAIR"
        else:
            rating = "[ERROR] NEEDS IMPROVEMENT"
        print(f"|  Rating: {rating:<48s}|")
    else:
        print("|  [WARN] No H5 dataset available for offline evaluation.   |")
        print("|  Set H5_PATH env var to enable ADE/FDE metrics.       |")
    print("+---------------------------------------------------------+")

    # -- Section 3: Online Metrics (from CARLA) --
    print("\n+---------------------------------------------------------+")
    print("|  [ONLINE] ONLINE METRICS (from CARLA simulation)             |")
    print("+---------------------------------------------------------+")
    if online_metrics:
        print(f"|  Route Completion:      {online_metrics.get('route', 'N/A'):<31s}|")
        print(f"|  Min TTC:               {online_metrics.get('ttc', 'N/A'):<31s}|")
        print(f"|  Total Frames:          {online_metrics.get('frames', 'N/A'):<31s}|")
    else:
        print("|  Route Completion:  N/A (Run --online to collect)      |")
        print("|  Min TTC:           N/A (Run --online to collect)      |")
        print("|                                                         |")
        print("|  To collect:                                            |")
        print("|    python scripts/evaluate_cnn_metrics.py --online      |")
    print("+---------------------------------------------------------+")

    # -- Summary Table --
    print("\n+--------------------------------------------------+")
    print("|           METRICS SUMMARY TABLE                  |")
    print("+-----------------------+--------------------------+")
    print("| Metric                | Value                    |")
    print("+-----------------------+--------------------------+")
    if ade_fde:
        print(f"| ADE                   | {ade_fde['ade_m']:>8.4f} m              |")
        print(f"| FDE                   | {ade_fde['fde_m']:>8.4f} m              |")
    else:
        print("| ADE                   | N/A (no dataset)         |")
        print("| FDE                   | N/A (no dataset)         |")
    print(f"| FPS (batch=1)         | {benchmark['fps']:>8.1f} FPS            |")
    print(f"| Latency (batch=1)     | {benchmark['latency_ms']:>8.2f} ms             |")
    
    if online_metrics:
        print(f"| Route Completion      | {online_metrics.get('route', 'N/A'):<24s} |")
        print(f"| TTC (Time-To-Collision| {online_metrics.get('ttc', 'N/A'):<24s} |")
    else:
        print("| Route Completion      | Run --online to collect  |")
        print("| TTC (Time-To-Collision| Run --online to collect  |")
    print("+-----------------------+--------------------------+")
    print()


# ============================================================================
# DOCX EXPORT
# ============================================================================
def _set_cell_bg(cell, hex_color: str):
    """Helper: set table cell background color via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_cell_text(cell, text: str, bold: bool = False,
                   font_size: int = 10, color: str | None = None,
                   align: str = "left"):
    """Helper: write text into a table cell with formatting."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def export_results_to_docx(
    benchmark: dict,
    ade_fde: dict | None,
    online_metrics: dict | None,
    output_dir: Path | None = None,
) -> Path | None:
    """
    Export evaluation results to a formatted .docx file under <output_dir>/
    (defaults to  <PROJECT_ROOT>/data/).

    Returns the Path to the written file, or None if python-docx is missing.
    """
    if not _DOCX_AVAILABLE:
        print(
            "\n[WARN] python-docx not installed — skipping .docx export.\n"
            "       Run:  pip install python-docx"
        )
        return None

    # ── Output directory ────────────────────────────────────────────────────
    if output_dir is None:
        output_dir = PROJECT_ROOT / "data"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = time.strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"evaluation_results_{timestamp}.docx"

    # ── Document setup ──────────────────────────────────────────────────────
    doc = Document()

    # Page margins (2 cm all sides)
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default font ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Title block ─────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("CIL WAYPOINT MODEL — COMPREHENSIVE EVALUATION REPORT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()  # spacer

    # ── Helper: section heading ──────────────────────────────────────────────
    def add_section_heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        # Bottom border for heading line
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # ── Helper: two-column metrics table ────────────────────────────────────
    HEADER_BG  = "1F497D"   # dark blue header rows
    ALT_BG     = "EEF3FA"   # light blue alternate rows
    WHITE_BG   = "FFFFFF"
    HEADER_FG  = "FFFFFF"

    def add_metrics_table(rows_data: list[tuple[str, str]]):
        """rows_data: list of (label, value) tuples."""
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"

        # Column widths (approx 60/40 split)
        for row_idx, (label, value) in enumerate(rows_data):
            row = tbl.add_row()
            bg = ALT_BG if row_idx % 2 == 0 else WHITE_BG
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[1], bg)

            row.cells[0].width = Cm(8)
            row.cells[1].width = Cm(6)

            _add_cell_text(row.cells[0], label, bold=True,  font_size=10)
            _add_cell_text(row.cells[1], value, bold=False, font_size=10)

        doc.add_paragraph()  # spacer after table

    # ── Helper: header row table (used for summary) ──────────────────────────
    def add_summary_table(headers: list[str], rows_data: list[list[str]]):
        tbl = doc.add_table(rows=0, cols=len(headers))
        tbl.style = "Table Grid"

        # Header row
        hdr_row = tbl.add_row()
        for i, h in enumerate(headers):
            _set_cell_bg(hdr_row.cells[i], HEADER_BG)
            _add_cell_text(hdr_row.cells[i], h, bold=True, font_size=10,
                           color=HEADER_FG, align="center")

        # Data rows
        for row_idx, row_data in enumerate(rows_data):
            row = tbl.add_row()
            bg = ALT_BG if row_idx % 2 == 0 else WHITE_BG
            for i, val in enumerate(row_data):
                _set_cell_bg(row.cells[i], bg)
                _add_cell_text(row.cells[i], val, font_size=10, align="center")

        doc.add_paragraph()

    # ── Section 1: Inference Performance ────────────────────────────────────
    add_section_heading("1. Inference Performance")

    fps      = benchmark.get("fps", 0.0)
    lat_ms   = benchmark.get("latency_ms", 0.0)
    lat_std  = benchmark.get("latency_std_ms", 0.0)
    batch_fp = benchmark.get("batch_fps", {})

    inf_rows = [
        ("FPS (batch=1)",     f"{fps:.1f} FPS"),
        ("Latency (batch=1)", f"{lat_ms:.2f} ± {lat_std:.2f} ms"),
    ]
    for bs, b_fps in batch_fp.items():
        if b_fps > 0:
            ms_per = 1000.0 / b_fps
            inf_rows.append((f"Throughput batch={bs}", f"{b_fps:.1f} FPS  ({ms_per:.2f} ms/sample)"))

    add_metrics_table(inf_rows)

    # ── Section 2: Trajectory Accuracy (Offline) ─────────────────────────────
    add_section_heading("2. Trajectory Accuracy (Offline)")

    if ade_fde:
        ade     = ade_fde["ade_m"]
        fde     = ade_fde["fde_m"]
        samples = ade_fde["total_samples"]
        if ade < 1.0 and fde < 1.5:
            rating = "EXCELLENT"
        elif ade < 2.0 and fde < 3.0:
            rating = "GOOD"
        elif ade < 3.0 and fde < 5.0:
            rating = "FAIR"
        else:
            rating = "NEEDS IMPROVEMENT"

        acc_rows = [
            ("ADE (Avg Displacement Error)",   f"{ade:.4f} m"),
            ("FDE (Final Displacement Error)", f"{fde:.4f} m"),
            ("Validation Samples",             f"{samples:,}"),
            ("Overall Rating",                 rating),
        ]
    else:
        acc_rows = [
            ("ADE", "N/A — set H5_PATH env var to enable"),
            ("FDE", "N/A — set H5_PATH env var to enable"),
        ]
    add_metrics_table(acc_rows)

    # ── Section 3: Online Metrics (CARLA) ───────────────────────────────────
    add_section_heading("3. Online Metrics (CARLA Simulation)")

    if online_metrics:
        online_rows = [
            ("Route Completion", online_metrics.get("route",  "N/A")),
            ("Min TTC (Time-To-Collision)", online_metrics.get("ttc", "N/A")),
            ("Total Frames",     online_metrics.get("frames", "N/A")),
        ]
    else:
        online_rows = [
            ("Route Completion",            "N/A — run with --online flag to collect"),
            ("Min TTC (Time-To-Collision)", "N/A — run with --online flag to collect"),
            ("Total Frames",                "N/A — run with --online flag to collect"),
        ]
        # Explanation note
        note = doc.add_paragraph()
        note_run = note.add_run(
            "Note: Online metrics require a running CARLA server. "
            "Re-run with:  python scripts/evaluate_cnn_metrics.py --online"
        )
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    add_metrics_table(online_rows)

    # ── Section 4: Summary Table ─────────────────────────────────────────────
    add_section_heading("4. Summary Table")

    summary_data = []
    summary_data.append(["ADE", f"{ade_fde['ade_m']:.4f} m" if ade_fde else "N/A"])
    summary_data.append(["FDE", f"{ade_fde['fde_m']:.4f} m" if ade_fde else "N/A"])
    summary_data.append(["FPS (batch=1)", f"{fps:.1f} FPS"])
    summary_data.append(["Latency (batch=1)", f"{lat_ms:.2f} ms"])
    if online_metrics:
        summary_data.append(["Route Completion",    online_metrics.get("route", "N/A")])
        summary_data.append(["TTC (Time-To-Collision)", online_metrics.get("ttc", "N/A")])
    else:
        summary_data.append(["Route Completion",       "N/A (--online not run)"])
        summary_data.append(["TTC (Time-To-Collision)", "N/A (--online not run)"])

    add_summary_table(["Metric", "Value"], summary_data)

    # ── Footer note ──────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "CIL Waypoint Model Evaluation  |  Auto-generated by evaluate_cnn_metrics.py"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(str(docx_path))
    print(f"\n[DOCX] Results exported to: {docx_path}")
    return docx_path


# ============================================================================
# MAIN
# ============================================================================
def main():
    parser = argparse.ArgumentParser(description="CIL Waypoint Model Evaluation")
    parser.add_argument("--benchmark-only", action="store_true",
                        help="Only run FPS/Latency benchmark (no dataset needed)")
    parser.add_argument("--online", action="store_true",
                        help="Launch CARLA simulation to collect online metrics (TTC, Route Completion)")
    parser.add_argument("--online-ticks", type=int, default=1000,
                        help="Number of CARLA ticks for online evaluation (default: 1000)")
    parser.add_argument("--online-timeout", type=int, default=300,
                        help="Timeout in seconds for online evaluation (default: 300)")
    parser.add_argument("--carla-host", default="127.0.0.1",
                        help="CARLA server host (default: 127.0.0.1)")
    parser.add_argument("--carla-port", type=int, default=2000,
                        help="CARLA server port (default: 2000)")
    parser.add_argument("--device", default="auto",
                        help="Device: 'cuda', 'cpu', or 'auto' (default)")
    parser.add_argument(
        "--csv",
        default=None,
        help=(
            "Path to a combined driving_log.csv covering all towns. "
            "The script will split it into per-town CSVs automatically. "
            "Overrides CSV_ROOT env var."
        ),
    )
    parser.add_argument(
        "--csv-root",
        default=None,
        help=(
            "Path to directory already structured as <root>/Town01/driving_log.csv, "
            "<root>/Town02/driving_log.csv, ... Overrides --csv and CSV_ROOT env var."
        ),
    )
    args = parser.parse_args()

    print("=" * 60)
    print(" CIL WAYPOINT MODEL - COMPREHENSIVE METRIC EVALUATION")
    print("=" * 60)

    # Device
    if args.device == "auto":
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    else:
        device = torch.device(args.device)
    print(f"Device: {device}")
    if device.type == "cuda":
        print(f"  GPU: {torch.cuda.get_device_name(0)}")
        print(f"  VRAM: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.1f} GB")

    # Load model
    model = load_model(device)

    # Always run benchmark
    benchmark = benchmark_inference(model, device)

    # Run ADE/FDE if dataset is available (and not --benchmark-only)
    ade_fde = None
    if not args.benchmark_only:
        h5_path = os.environ.get("H5_PATH",
                                  "/kaggle/input/datasets/yudtrann/dataset-carlav3/carla_images_drive.h5")
        # Resolve csv_root_hint — priority: --csv-root > --csv > CSV_ROOT env > train_params.yaml
        if args.csv_root:
            csv_root_hint = args.csv_root
        elif args.csv:
            csv_root_hint = args.csv  # evaluate_ade_fde detects it's a file and splits
        else:
            try:
                import yaml
                config_path = PROJECT_ROOT / "configs" / "train_params.yaml"
                with open(config_path, "r", encoding="utf-8") as f:
                    config = yaml.safe_load(f)
                csv_root_hint = os.environ.get("CSV_ROOT", config.get("data_root", None))
            except Exception:
                csv_root_hint = None

        ade_fde = evaluate_ade_fde(model, device, h5_path, csv_root_hint)

    # Run online metrics (from CARLA simulation)
    online_metrics = None
    if getattr(args, "online", False):
        import subprocess
        import re

        print("\n" + "=" * 60)
        print(" [ONLINE] RUNNING CARLA SIMULATION FOR ONLINE METRICS")
        print("=" * 60)

        # ── Parse configuration from run_cil.ps1 ─────────────────────────
        # This mirrors *exactly* what run_cil.ps1 does:
        #   • reads all default paths (CarlaRoot, YoloModel, Config, Model, Agent, Device)
        #   • sets CARLA_ROOT / CARLA_PYTHONAPI / PYTHONPATH env vars
        #   • calls run_agents.py with --agent cil_yolo + all required flags
        ps1_path = PROJECT_ROOT / "scripts" / "run_cil.ps1"
        ps1_cfg  = _parse_ps1_defaults(ps1_path)

        carla_root       = ps1_cfg["carla_root"]                        # e.g. E:\Carla
        carla_pythonapi  = str(Path(carla_root) / "PythonAPI")
        carla_agents_dir = str(Path(carla_root) / "PythonAPI" / "carla")

        agent      = ps1_cfg["agent"]      # "cil_yolo"
        config     = ps1_cfg["config"]     # "configs/carla_env.yaml"
        device     = ps1_cfg["device"]     # "auto"

        # Resolve model paths (same logic as run_cil.ps1)
        raw_model = ps1_cfg["model"]
        cil_model_path = (
            raw_model if Path(raw_model).is_absolute()
            else str(PROJECT_ROOT / raw_model)
        )
        raw_yolo = ps1_cfg["yolo_model"]
        yolo_model_path = (
            raw_yolo if Path(raw_yolo).is_absolute()
            else str(PROJECT_ROOT / raw_yolo)
        )

        # Resolve config path
        config_path = (
            config if Path(config).is_absolute()
            else str(PROJECT_ROOT / config)
        )

        # Locate run_agents.py (at project root, same as run_cil.ps1 does)
        run_agents_path = PROJECT_ROOT / "run_agents.py"
        if not run_agents_path.exists():
            run_agents_path = PROJECT_ROOT / "scripts" / "run_agents.py"

        # Build command — identical to run_cil.ps1's $runnerArgs block
        cmd = [
            sys.executable, str(run_agents_path),
            "--agent",          agent,
            "--config",         config_path,
            "--cil-model-path", cil_model_path,
            "--yolo-model-path", yolo_model_path,
            "--device",         device,
            "--ticks",          str(args.online_ticks),
            "--eval-online",    # ← guarantees teardown() prints the metrics block
            "--host",           args.carla_host,
            "--port",           str(args.carla_port),
        ]

        # Set environment variables matching run_cil.ps1's env setup block
        env = os.environ.copy()
        env["CARLA_ROOT"]       = carla_root
        env["CARLA_PYTHONAPI"]  = carla_pythonapi
        existing_pp = env.get("PYTHONPATH", "")
        env["PYTHONPATH"] = (carla_agents_dir + ";" + existing_pp).rstrip(";")

        print(f"  Agent      : {agent}")
        print(f"  Config     : {config_path}")
        print(f"  CIL model  : {cil_model_path}")
        print(f"  YOLO model : {yolo_model_path}")
        print(f"  CARLA root : {carla_root}")
        print(f"  Ticks      : {args.online_ticks}  |  Timeout: {args.online_timeout}s")
        print(f"  Host:Port  : {args.carla_host}:{args.carla_port}")
        print("  Please wait...\n")

        try:
            result = subprocess.run(
                cmd,
                cwd=str(PROJECT_ROOT),
                env=env,
                capture_output=True,
                text=True,
                timeout=args.online_timeout,
                encoding="utf-8",
                errors="replace",
            )

            output = result.stdout + "\n" + result.stderr

            # Patterns match the teardown() block in run_agents.py:
            #   " Total Frames Ran:       {n}"
            #   " Minimum TTC observed:   {v:.2f} s"  or  "... N/A (No vehicles ahead)"
            #   " Route Completion:       {v:.1f}%"
            frames_match = re.search(r"Total Frames Ran:\s*(.+)",     output)
            ttc_match    = re.search(r"Minimum TTC observed:\s*(.+)", output)
            route_match  = re.search(r"Route Completion:\s*(.+)",     output)

            online_metrics = {
                "frames": frames_match.group(1).strip() if frames_match else "N/A",
                "ttc":    ttc_match.group(1).strip()    if ttc_match    else "N/A",
                "route":  route_match.group(1).strip()  if route_match  else "N/A",
            }

            if all(v == "N/A" for v in online_metrics.values()):
                print("  [WARN] Could not parse any online metrics from subprocess output.")
                print("  Possible causes:")
                print(f"    1. CARLA server not running on {args.carla_host}:{args.carla_port}")
                print("    2. run_agents.py crashed before teardown() printed metrics")
                print(f"    3. Return code: {result.returncode}")
                lines = output.strip().splitlines()
                if lines:
                    print("\n  --- Captured output (last 40 lines) ---")
                    for ln in lines[-40:]:
                        print(f"    {ln}")
                    print("  --- End ---")
                else:
                    print("  [WARN] No output was captured at all.")
            else:
                print("  [OK] Online evaluation finished.")

        except subprocess.TimeoutExpired:
            print(f"  [ERROR] Timed out after {args.online_timeout}s.")
            online_metrics = {"frames": "Timeout", "route": "Timeout", "ttc": "Timeout"}
        except Exception as e:
            print(f"  [ERROR] Failed to run online evaluation: {e}")

    # Print results
    print_results(benchmark, ade_fde, online_metrics)

    # Export results to .docx in the data/ folder
    export_results_to_docx(
        benchmark=benchmark,
        ade_fde=ade_fde,
        online_metrics=online_metrics,
        output_dir=PROJECT_ROOT / "data",
    )


if __name__ == "__main__":
    main()
